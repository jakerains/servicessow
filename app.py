import streamlit as st
import boto3
import json
import os
import time
import requests
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from dotenv import load_dotenv
from pydub import AudioSegment
import multiprocessing
import concurrent.futures

# Load environment variables
load_dotenv()

def load_questions():
    """Load questions from Questions.json file"""
    try:
        with open('Questions.json', 'r') as f:
            return json.load(f)
    except Exception as e:
        st.error(f"Error loading questions: {str(e)}")
        return None

# Initialize session state
if 'aws_credentials' not in st.session_state:
    # Try environment variables first
    aws_credentials = {
        'aws_access_key_id': os.getenv('AWS_ACCESS_KEY_ID'),
        'aws_secret_access_key': os.getenv('AWS_SECRET_ACCESS_KEY'),
        'region_name': os.getenv('AWS_REGION', 'us-east-1'),
        'bucket_name': os.getenv('AWS_BUCKET_NAME')
    }
    
    # If not in environment variables, try Streamlit secrets
    if not all(aws_credentials.values()):
        try:
            aws_credentials = {
                'aws_access_key_id': st.secrets["aws"]["aws_access_key_id"],
                'aws_secret_access_key': st.secrets["aws"]["aws_secret_access_key"],
                'region_name': st.secrets["aws"]["aws_region"],
                'bucket_name': st.secrets["aws"]["aws_bucket_name"]
            }
        except Exception as e:
            st.error(f"Error loading AWS credentials: {str(e)}")
            aws_credentials = None
    
    st.session_state['aws_credentials'] = aws_credentials

if 'questions' not in st.session_state:
    st.session_state['questions'] = load_questions()

if 'results' not in st.session_state:
    st.session_state['results'] = None

if 'transcription' not in st.session_state:
    st.session_state['transcription'] = None

def transcribe_audio(file, aws_credentials):
    """Transcribe audio file using AWS Transcribe with detailed progress tracking"""
    try:
        transcribe = boto3.client(
            'transcribe',
            aws_access_key_id=aws_credentials['aws_access_key_id'],
            aws_secret_access_key=aws_credentials['aws_secret_access_key'],
            region_name=aws_credentials['region_name']
        )
        
        s3 = boto3.client('s3', **{k: aws_credentials[k] for k in ['aws_access_key_id', 'aws_secret_access_key', 'region_name']})
        bucket_name = aws_credentials['bucket_name']
        file_name = f"audio_{int(time.time())}.mp3"
        
        # Create progress indicators
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Step 1: Upload to S3
        status_text.text("Uploading audio file to S3...")
        progress_bar.progress(10)
        s3.upload_fileobj(file, bucket_name, file_name)
        progress_bar.progress(20)
        
        # Step 2: Start transcription
        status_text.text("Starting transcription job...")
        job_name = f"transcription_{int(time.time())}"
        transcribe.start_transcription_job(
            TranscriptionJobName=job_name,
            Media={'MediaFileUri': f"s3://{bucket_name}/{file_name}"},
            MediaFormat='mp3',
            LanguageCode='en-US',
            Settings={
                'ShowSpeakerLabels': True,
                'MaxSpeakerLabels': 2,
            }
        )
        progress_bar.progress(30)
        
        # Step 3: Monitor progress
        while True:
            status = transcribe.get_transcription_job(TranscriptionJobName=job_name)
            job_status = status['TranscriptionJob']['TranscriptionJobStatus']
            
            if job_status == 'IN_PROGRESS':
                # Get progress details
                if 'Progress' in status['TranscriptionJob']:
                    progress = status['TranscriptionJob']['Progress']
                    progress_percent = int(30 + (progress * 60))  # Scale from 30% to 90%
                    progress_bar.progress(progress_percent)
                    status_text.text(f"Transcribing audio... {progress_percent}% complete")
                
            elif job_status == 'COMPLETED':
                status_text.text("Retrieving transcription results...")
                progress_bar.progress(90)
                response = requests.get(status['TranscriptionJob']['Transcript']['TranscriptFileUri'])
                transcription = response.json()['results']['transcripts'][0]['transcript']
                
                # Clean up S3
                status_text.text("Cleaning up temporary files...")
                s3.delete_object(Bucket=bucket_name, Key=file_name)
                
                progress_bar.progress(100)
                status_text.text("Transcription complete!")
                time.sleep(1)  # Let user see completion
                status_text.empty()
                progress_bar.empty()
                
                return transcription
                
            elif job_status == 'FAILED':
                error_reason = status['TranscriptionJob'].get('FailureReason', 'Unknown error')
                status_text.text(f"Transcription failed: {error_reason}")
                progress_bar.empty()
                raise Exception(f"Transcription failed: {error_reason}")
            
            time.sleep(5)  # Poll every 5 seconds
            
    except Exception as e:
        status_text.text(f"Error: {str(e)}")
        progress_bar.empty()
        st.error(f"Error in transcription: {str(e)}")
        return None

def start_transcription_job(transcribe_client, job_name, s3_uri, aws_credentials):
    """Helper function to manage a single transcription job"""
    try:
        transcribe_client.start_transcription_job(
            TranscriptionJobName=job_name,
            Media={'MediaFileUri': s3_uri},
            MediaFormat='mp3',
            LanguageCode='en-US',
            Settings={
                'ShowSpeakerLabels': True,
                'MaxSpeakerLabels': 2,
            }
        )
        
        while True:
            status = transcribe_client.get_transcription_job(TranscriptionJobName=job_name)
            job_status = status['TranscriptionJob']['TranscriptionJobStatus']
            
            if job_status in ['COMPLETED', 'FAILED']:
                break
            time.sleep(5)
        
        if job_status == 'COMPLETED':
            response = requests.get(status['TranscriptionJob']['Transcript']['TranscriptFileUri'])
            return response.json()['results']['transcripts'][0]['transcript']
        else:
            raise Exception(f"Transcription failed: {status['TranscriptionJob'].get('FailureReason')}")
            
    except Exception as e:
        raise Exception(f"Error in transcription job: {str(e)}")

def process_with_bedrock(text, questions, aws_credentials):
    """Process questions in parallel batches"""
    import concurrent.futures
    
    def process_batch(batch):
        results = []
        for category in batch:
            # Process each category
            category_results = {
                'category': category['category'],
                'answers': []
            }
            
            for question in category['questions']:
                if isinstance(question, dict):
                    question_text = question['text']
                    instruction = question.get('instruction', '')
                else:
                    question_text = question
                    instruction = ''
                
                prompt = f"""Based on the following text, please answer this question: {question_text}

                {f'Additional instruction: {instruction}' if instruction else ''}

                Text: {text}

                Please provide a clear, detailed, and professional answer based only on the information provided in the text. 
                If the information is not available in the text, please indicate that clearly."""

                try:
                    response = bedrock_runtime.invoke_model(
                        modelId='us.anthropic.claude-3-5-haiku-20241022-v1:0',
                        body=json.dumps({
                            "anthropic_version": "bedrock-2023-05-31",
                            "max_tokens": 1000,
                            "messages": [
                                {
                                    "role": "user",
                                    "content": prompt
                                }
                            ]
                        })
                    )
                    
                    response_body = json.loads(response['body'].read())
                    answer = {
                        "question": question_text,
                        "answer": response_body['content'][0]['text']
                    }
                    category_results["answers"].append(answer)
                    
                except Exception as e:
                    st.error(f"Error processing question '{question_text}': {str(e)}")
                    continue
            
            results.append(category_results)
        return results
    
    try:
        bedrock_runtime = boto3.client(
            service_name='bedrock-runtime',
            aws_access_key_id=aws_credentials['aws_access_key_id'],
            aws_secret_access_key=aws_credentials['aws_secret_access_key'],
            region_name=aws_credentials['region_name']
        )
        
        # Split questions into batches
        batch_size = 3  # Process 3 categories at a time
        batches = [questions["project_questions"][i:i + batch_size] for i in range(0, len(questions["project_questions"]), batch_size)]
        
        # Process batches in parallel
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            batch_results = list(executor.map(process_batch, batches))
        
        # Combine results
        final_results = []
        for batch in batch_results:
            final_results.extend(batch)
        
        return final_results
        
    except Exception as e:
        st.error(f"Error initializing Bedrock: {str(e)}")
        return None

def generate_text_results(results):
    """Generate text format results"""
    text_output = ""
    for category in results:
        text_output += f"\n{category['category']}\n{'='*len(category['category'])}\n\n"
        for qa in category['answers']:
            text_output += f"Q: {qa['question']}\nA: {qa['answer']}\n\n"
    return text_output

def generate_pdf_results(results):
    """Generate PDF format results"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    for category in results:
        elements.append(Paragraph(category['category'], styles['Heading1']))
        elements.append(Spacer(1, 12))
        for qa in category['answers']:
            elements.append(Paragraph(f"Q: {qa['question']}", styles['Heading2']))
            elements.append(Paragraph(f"A: {qa['answer']}", styles['Normal']))
            elements.append(Spacer(1, 12))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_docx_results(results):
    """Generate DOCX format results"""
    doc = Document()
    
    # Define styles using names instead of IDs
    for category in results:
        # Use built-in heading style names
        heading = doc.add_paragraph(category['category'])
        heading.style = 'Heading 1'
        
        for qa in category['answers']:
            # Question with Heading 2
            question = doc.add_paragraph(f"Q: {qa['question']}")
            question.style = 'Heading 2'
            
            # Answer with Normal style
            answer = doc.add_paragraph(f"A: {qa['answer']}")
            answer.style = 'Normal'
            
            # Add spacing
            doc.add_paragraph()
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def edit_questions():
    """Edit questions in the Questions.json file"""
    questions = st.session_state.get('questions', load_questions())
    
    if questions:
        edited_questions = {"project_questions": []}
        
        st.subheader("Edit Questions")
        st.write("Add, edit, or remove categories and questions below.")
        
        # Add new category button
        if st.button("Add New Category"):
            questions["project_questions"].append({
                "category": "New Category",
                "questions": ["New Question"]
            })
        
        # Edit existing categories and questions
        for cat_idx, category in enumerate(questions["project_questions"]):
            st.markdown("---")
            col1, col2 = st.columns([3, 1])
            
            with col1:
                new_category = st.text_input(
                    f"Category {cat_idx + 1}",
                    value=category["category"],
                    key=f"cat_{cat_idx}"
                )
            
            with col2:
                if st.button("Delete Category", key=f"del_cat_{cat_idx}"):
                    continue
            
            category_questions = []
            for q_idx, question in enumerate(category["questions"]):
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    if isinstance(question, dict):
                        question_text = st.text_input(
                            f"Question {cat_idx + 1}.{q_idx + 1}",
                            value=question["text"],
                            key=f"q_{cat_idx}_{q_idx}"
                        )
                        instruction = st.text_input(
                            "Instruction (optional)",
                            value=question.get("instruction", ""),
                            key=f"i_{cat_idx}_{q_idx}"
                        )
                        if instruction:
                            category_questions.append({
                                "text": question_text,
                                "instruction": instruction
                            })
                        else:
                            category_questions.append(question_text)
                    else:
                        question_text = st.text_input(
                            f"Question {cat_idx + 1}.{q_idx + 1}",
                            value=question,
                            key=f"q_{cat_idx}_{q_idx}"
                        )
                        category_questions.append(question_text)
                
                with col2:
                    if st.button("Delete Question", key=f"del_q_{cat_idx}_{q_idx}"):
                        continue
            
            # Add new question button for this category
            if st.button("Add Question", key=f"add_q_{cat_idx}"):
                category_questions.append("New Question")
            
            if category_questions:  # Only add category if it has questions
                edited_questions["project_questions"].append({
                    "category": new_category,
                    "questions": category_questions
                })
        
        # Save changes button
        if st.button("Save Changes"):
            try:
                with open('Questions.json', 'w') as f:
                    json.dump(edited_questions, f, indent=4)
                st.session_state['questions'] = edited_questions
                st.success("Questions saved successfully!")
            except Exception as e:
                st.error(f"Error saving questions: {str(e)}")

def main():
    st.title("Sterling Services: S.O.W. Generator")
    
    tab1, tab2 = st.tabs(["Generate SOW", "Edit Questions"])
    
    with tab1:
        if not st.session_state.get('aws_credentials'):
            st.error("AWS credentials not found. Please check your environment variables or secrets.toml file.")
            return
        
        uploaded_file = st.file_uploader("Upload an audio or text file", type=['mp3', 'wav', 'txt', 'pdf', 'docx'])
        
        if uploaded_file is not None:
            # Process the file
            if uploaded_file.type.startswith('audio/'):
                if 'transcription' not in st.session_state or not st.session_state['transcription']:
                    with st.spinner('Transcribing audio...'):
                        transcription = transcribe_audio(uploaded_file, st.session_state['aws_credentials'])
                        if transcription:
                            st.session_state['transcription'] = transcription
                            st.success("Audio transcribed successfully!")
                            
                            # Automatically start analysis after transcription
                            with st.spinner('Analyzing content with AWS Bedrock...'):
                                results = process_with_bedrock(
                                    st.session_state['transcription'],
                                    st.session_state['questions'],
                                    st.session_state['aws_credentials']
                                )
                                if results:
                                    st.session_state['results'] = results
                                    st.success("Analysis complete!")
                                    st.rerun()
            else:
                # Handle text files
                text_content = uploaded_file.read().decode('utf-8')
                st.session_state['transcription'] = text_content
                
                # Automatically start analysis for text files too
                with st.spinner('Analyzing content with AWS Bedrock...'):
                    results = process_with_bedrock(
                        st.session_state['transcription'],
                        st.session_state['questions'],
                        st.session_state['aws_credentials']
                    )
                    if results:
                        st.session_state['results'] = results
                        st.success("Analysis complete!")
                        st.rerun()
            
            # Show transcription preview and results
            if st.session_state.get('transcription'):
                st.write("### Content Preview:")
                st.write(st.session_state['transcription'][:500] + "..." if len(st.session_state['transcription']) > 500 else st.session_state['transcription'])
                
                # Display results if available
                if st.session_state.get('results'):
                    st.write("### Analysis Results")
                    for category in st.session_state['results']:
                        with st.expander(f"📋 {category['category']}", expanded=True):
                            for qa in category['answers']:
                                st.markdown(f"**Q:** {qa['question']}")
                                st.markdown(f"**A:** {qa['answer']}")
                                st.markdown("---")
                    
                    # Download buttons
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        text_output = generate_text_results(st.session_state['results'])
                        st.download_button(
                            "Download TXT",
                            text_output,
                            "results.txt",
                            "text/plain"
                        )
                    with col2:
                        pdf_output = generate_pdf_results(st.session_state['results'])
                        st.download_button(
                            "Download PDF",
                            pdf_output,
                            "results.pdf",
                            "application/pdf"
                        )
                    with col3:
                        docx_output = generate_docx_results(st.session_state['results'])
                        st.download_button(
                            "Download DOCX",
                            docx_output,
                            "results.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
    
    with tab2:
        edit_questions()

if __name__ == "__main__":
    main()