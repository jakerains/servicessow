import streamlit as st
import os
import json
from groq import Groq
from groq.types.model_list import ModelList
from groq.types.model import Model
import pandas as pd
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from io import BytesIO
import docx2txt
import PyPDF2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import markdown
import base64
import mimetypes
from bs4 import BeautifulSoup
import tempfile
import requests
import random

# Define the version number
VERSION = "2.0.1"  # Updated version number

# Set page config at the very top, after imports
st.set_page_config(page_title="Sterling Services: S.O.W. Generator (Groq Version)", page_icon="📄")

# Initialize session state
if 'api_key' not in st.session_state:
    st.session_state['api_key'] = ''
if 'api_key_set' not in st.session_state:
    st.session_state['api_key_set'] = False
if 'groq_client' not in st.session_state:
    st.session_state['groq_client'] = None

def init_groq_client(api_key):
    try:
        client = Groq(api_key=api_key)
        # Test if the client has the 'models' attribute
        client.models.list()
        st.success("Groq client initialized successfully")
        return client
    except Exception as e:
        st.error(f"Failed to initialize Groq client: {str(e)}")
        return None

def fetch_groq_models():
    if st.session_state['groq_client'] is None:
        st.error("Groq client is not initialized. Cannot fetch models.")
        return []

    try:
        models = st.session_state['groq_client'].models.list()
        
        if isinstance(models, ModelList) and hasattr(models, 'data'):
            chat_models = [model.id for model in models.data if isinstance(model, Model) and not model.id.startswith('whisper')]
        else:
            st.error(f"Unexpected model list structure: {type(models)}")
            return []
        
        print(f"Extracted chat models: {chat_models}")  # This will show us the final list of models
        return chat_models
    except Exception as e:
        st.error(f"Failed to fetch Groq models: {str(e)}")
        return []

def transcribe_audio(file):
    api_key = st.session_state['api_key']
    if not api_key:
        st.error("Groq API key not found. Please set it in your environment variables or Streamlit secrets.")
        return None

    progress_bar = st.progress(0)
    status_text = st.empty()

    try:
        status_text.text("Preparing audio for transcription...")
        progress_bar.progress(30)

        url = "https://api.groq.com/openai/v1/audio/transcriptions"
        headers = {
            "Authorization": f"Bearer {api_key}"
        }
        files = {
            "file": file,
            "model": (None, "distil-whisper-large-v3-en"),
            "response_format": (None, "text")
        }

        response = requests.post(url, headers=headers, files=files)
        response.raise_for_status()  # Raise an error for bad status codes

        transcription = response.text
        print(f"Transcription type: {type(transcription)}")
        print(f"Transcription content: {transcription[:100]}...")  # Print first 100 characters

        progress_bar.progress(100)
        status_text.text("Transcription complete!")

        return transcription

    except requests.exceptions.RequestException as e:
        st.error(f"An error occurred during transcription: {str(e)}")
        return None

    finally:
        progress_bar.empty()
        status_text.empty()

def detect_file_type(file):
    # Get the file extension
    file_extension = os.path.splitext(file.name)[1].lower()
    
    # Define file type based on extension
    audio_extensions = ['.mp3', '.wav', '.m4a', '.aac', '.flac']
    text_extensions = ['.txt', '.rtf', '.doc', '.docx', '.pdf', '.odt', '.md']
    
    if file_extension in audio_extensions:
        return "Audio"
    elif file_extension in text_extensions:
        return "Text"
    else:
        # If the extension is not recognized, try to guess the MIME type
        mime_type, _ = mimetypes.guess_type(file.name)
        if mime_type:
            if mime_type.startswith('audio'):
                return "Audio"
            elif mime_type.startswith('text') or mime_type == 'application/pdf':
                return "Text"
    
    # If we can't determine the type, return None
    return None

def process_file(file):
    file_extension = os.path.splitext(file.name)[1].lower()
    
    if file_extension in ['.txt', '.md']:
        return file.getvalue().decode('utf-8')
    elif file_extension == '.pdf':
        pdf_reader = PyPDF2.PdfReader(file)
        return ' '.join([page.extract_text() for page in pdf_reader.pages])
    elif file_extension in ['.doc', '.docx']:
        doc = Document(file)
        return ' '.join([para.text for para in doc.paragraphs])
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def process_transcription(text, questions, model_name):
    client = st.session_state['groq_client']
    
    for category in questions["project_questions"]:
        category_results = {"category": category["category"], "answers": []}
        for question in category["questions"]:
            if isinstance(question, dict):
                question_text = question['text']
            else:
                question_text = question
            prompt = f"Based on the following text, answer this question: {question_text}\n\nText: {text}"
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=150
            )
            answer = {
                "question": question_text,
                "answer": response.choices[0].message.content.strip()
            }
            category_results["answers"].append(answer)
            yield category_results["category"], answer

    return

def load_questions():
    # Define the path to your questions JSON file
    questions_file = "questions.json"
    
    try:
        # Open and read the JSON file
        with open(questions_file, "r") as file:
            questions = json.load(file)
        return questions
    except FileNotFoundError:
        st.error(f"Questions file '{questions_file}' not found.")
        return None
    except json.JSONDecodeError:
        st.error(f"Error decoding JSON from '{questions_file}'. Please check the file format.")
        return None

def main():
    st.title("Sterling Services: S.O.W. Generator (Groq Version)")
    
    # Add New Project button at the top
    if st.button("New Project"):
        # Reset all relevant session state variables
        st.session_state['uploaded_file'] = None
        st.session_state['analysis_results'] = None
        st.session_state['transcription'] = None
        st.session_state['file_uploader_key'] = str(random.randint(1000, 9999))  # Add this line
        st.rerun()
    
    # Configuration section
    st.subheader("Configuration")
    
    # Initialize session state
    if 'api_key' not in st.session_state:
        st.session_state['api_key'] = ''
    if 'api_key_set' not in st.session_state:
        st.session_state['api_key_set'] = False
    if 'model_name' not in st.session_state:
        st.session_state['model_name'] = ''
    if 'questions' not in st.session_state:
        st.session_state['questions'] = load_questions()
    if 'show_questions' not in st.session_state:
        st.session_state['show_questions'] = False
    if 'analysis_results' not in st.session_state:
        st.session_state['analysis_results'] = None
    if 'config_set' not in st.session_state:
        st.session_state['config_set'] = False
    if 'groq_models' not in st.session_state:
        st.session_state['groq_models'] = []

    # Only ask for the API key if it's not already set
    if not st.session_state['api_key_set']:
        api_key = st.text_input("Enter your Groq API key:", type="password", key="api_key_input")
        if api_key:
            st.session_state['api_key'] = api_key
            st.session_state['api_key_set'] = True
            st.session_state['groq_client'] = init_groq_client(api_key)
            st.success("API key saved for this session.")
            st.rerun()

    if st.session_state['api_key_set']:
        st.success("API key is set for this session.")

        # Fetch Groq models if not already fetched
        if 'groq_models' not in st.session_state or not st.session_state['groq_models']:
            st.session_state['groq_models'] = fetch_groq_models()

        # Model selection dropdown
        if st.session_state['groq_models']:
            default_model = 'llama-3.1-70b-versatile'  # Updated default model
            default_index = st.session_state['groq_models'].index(default_model) if default_model in st.session_state['groq_models'] else 0
            
            model_name = st.selectbox(
                "Select Groq model",
                options=st.session_state['groq_models'],
                index=default_index,
                key='model_selector'
            )
            
            if st.button("Set Configuration", key="set_config_button"):
                st.session_state['model_name'] = model_name
                st.session_state['config_set'] = True
                st.success(f"Configuration set: Using model {model_name}")
        else:
            st.error("Failed to fetch Groq models. Please check your API key and try again.")

        # Questions section
        st.subheader("Questions")
        if st.button("Hide Questions" if st.session_state['show_questions'] else "Show Questions", key="toggle_questions_button"):
            st.session_state['show_questions'] = not st.session_state['show_questions']
            st.rerun()

        # Display and customize questions
        if st.session_state['show_questions']:
            display_and_customize_questions()

        # File upload section
        st.subheader("File Upload")
        st.write("Upload your audio or text file")
        
        # Use a unique key for the file_uploader
        file_uploader_key = st.session_state.get('file_uploader_key', '0')
        file = st.file_uploader("Choose a file", type=["mp3", "wav", "m4a", "txt", "rtf", "doc", "docx", "pdf", "odt", "md"], key=file_uploader_key)
        
        if file is not None:
            st.session_state['uploaded_file'] = file
        elif 'uploaded_file' in st.session_state:
            file = st.session_state['uploaded_file']

        if file:
            file_type = detect_file_type(file)
            
            if file_type is None:
                st.error("Unsupported file type. Please upload an audio or text file.")
            else:
                st.write(f"File uploaded: {file.name} (Detected as: {file_type})")
                
                if st.button("Analyze", key="analyze_button"):
                    try:
                        if file_type == "Audio":
                            st.write(f"Processing audio file: {file.name}")
                            transcription = transcribe_audio(file)
                            if transcription is None:
                                st.error("Transcription failed. Please try again.")
                                return
                            st.success("Transcription complete!")
                            # Store the transcription in session state for later use
                            st.session_state['transcription'] = transcription
                        else:
                            st.write(f"Processing text file: {file.name}")
                            transcription = process_file(file)
                        
                        # Reset analysis results
                        st.session_state['analysis_results'] = None
                        
                        # Process and display results
                        process_and_display_results(transcription, st.session_state['questions'])
                    except Exception as e:
                        st.error(f"An error occurred while processing the file: {str(e)}")
                        st.write("Please try again or contact support if the issue persists.")
                
                # Add this condition to display results if they exist
                elif st.session_state['analysis_results'] is not None:
                    process_and_display_results(None, None)

    # Initialize session state for changelog visibility
    if 'show_changelog' not in st.session_state:
        st.session_state['show_changelog'] = False

    # Add version number and changelog at the bottom
    st.markdown("---")
    st.markdown(f"<p style='text-align: center; color: gray;'>Version {VERSION}</p>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: gray;'>Created by GenAI Jake</p>", unsafe_allow_html=True)
    
    # Toggle changelog visibility
    if st.button("Hide Changelog" if st.session_state['show_changelog'] else "View Changelog", key="toggle_changelog_button"):
        st.session_state['show_changelog'] = not st.session_state['show_changelog']
        st.rerun()

    # Display changelog if show_changelog is True
    if st.session_state['show_changelog']:
        display_changelog()

def process_and_display_results(transcription, questions):
    if st.session_state['analysis_results'] is None:
        with st.spinner(f"Analyzing content using {st.session_state['model_name']}..."):
            if isinstance(transcription, str):
                text_to_process = transcription
            else:
                st.error(f"Invalid transcription format: {type(transcription)}")
                st.write("Transcription content:", transcription)
                return

            results = []
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            total_questions = sum(len(category["questions"]) for category in questions["project_questions"])
            processed_questions = 0

            client = st.session_state['groq_client']

            for category in questions["project_questions"]:
                category_results = {"category": category["category"], "answers": []}
                for question in category["questions"]:
                    question_text = question['text'] if isinstance(question, dict) else question
                    prompt = f"Based on the following text, answer this question: {question_text}\n\nText: {text_to_process}"
                    response = client.chat.completions.create(
                        model=st.session_state['model_name'],
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=150
                    )
                    answer = {
                        "question": question_text,
                        "answer": response.choices[0].message.content.strip()
                    }
                    category_results["answers"].append(answer)
                    
                    processed_questions += 1
                    progress = processed_questions / total_questions
                    progress_bar.progress(progress)
                    status_text.text(f"Processing: {category['category']} - {question_text}")

                results.append(category_results)

            progress_bar.empty()
            status_text.empty()

        st.session_state['analysis_results'] = results
        st.session_state['transcription'] = text_to_process
    else:
        results = st.session_state['analysis_results']
    
    st.subheader("Analysis Results")
    
    # Generate results in different formats
    text_results = generate_text_results(results)
    pdf_results = generate_pdf_results(results)
    docx_results = generate_docx_results(results)
    
    # Create download buttons at the top
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.download_button(
            label="Download as TXT",
            data=text_results,
            file_name="analysis_results.txt",
            mime="text/plain",
            key="download_txt"
        )
    with col2:
        st.download_button(
            label="Download as PDF",
            data=pdf_results,
            file_name="analysis_results.pdf",
            mime="application/pdf",
            key="download_pdf"
        )
    with col3:
        st.download_button(
            label="Download as DOCX",
            data=docx_results,
            file_name="analysis_results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx"
        )
    with col4:
        if 'transcription' in st.session_state and st.session_state['transcription']:
            st.download_button(
                label="Download Transcript",
                data=st.session_state['transcription'],
                file_name="original_transcript.txt",
                mime="text/plain",
                key="download_transcript"
            )
    
    # Display results
    display_results(results)

def display_results(results):
    for category in results:
        st.subheader(category['category'])
        for qa in category['answers']:
            st.markdown(f"**Q: {qa['question']}**")
            st.write(f"A: {qa['answer']}")
            st.markdown("---")

def generate_text_results(results):
    text_output = ""
    for category in results:
        text_output += f"{category['category']}\n\n"
        for qa in category['answers']:
            text_output += f"Q: {qa['question']}\nA: {qa['answer']}\n\n"
        text_output += "\n"
    return text_output

def generate_pdf_results(results):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create custom styles
    styles.add(ParagraphStyle(name='Justify', parent=styles['BodyText'], alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name='Heading1Custom', parent=styles['Heading1'], fontSize=16, spaceAfter=12))
    styles.add(ParagraphStyle(name='Heading3Custom', parent=styles['Heading3'], fontSize=14, spaceAfter=8))
    
    story = []

    for category in results:
        story.append(Paragraph(category['category'], styles['Heading1Custom']))
        for qa in category['answers']:
            story.append(Paragraph(f"Q: {qa['question']}", styles['Heading3Custom']))
            story.append(Paragraph(f"A: {qa['answer']}", styles['Justify']))
            story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_docx_results(results):
    doc = Document()
    
    # Define styles
    styles = doc.styles
    heading1_style = styles['Heading 1']
    heading3_style = styles['Heading 3']
    normal_style = styles['Normal']

    for category in results:
        heading = doc.add_paragraph(category['category'])
        heading.style = heading1_style
        for qa in category['answers']:
            question = doc.add_paragraph(f"Q: {qa['question']}")
            question.style = heading3_style
            answer = doc.add_paragraph(f"A: {qa['answer']}")
            answer.style = normal_style
            doc.add_paragraph()  # Add a blank line

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def format_transcript(transcription):
    formatted_text = ""
    for segment in transcription['segments']:
        start_time = format_time(segment['start'])
        end_time = format_time(segment['end'])
        formatted_text += f"[{start_time} - {end_time}] {segment['text']}\n"
    return formatted_text

def format_time(seconds):
    minutes, seconds = divmod(int(seconds), 60)
    hours, minutes = divmod(minutes, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

def display_and_customize_questions():
    st.subheader("Review and Customize Questions")
    
    for category_index, category in enumerate(st.session_state['questions']["project_questions"]):
        st.markdown(f"### {category['category']}")
        for i, question in enumerate(category["questions"]):
            with st.container():
                col1, col2, col3 = st.columns([6, 1, 1])
                with col1:
                    q_text = question["text"] if isinstance(question, dict) else question
                    new_text = st.text_input(
                        f"Question {i+1}",
                        value=q_text,
                        key=f"q_{category_index}_{i}",
                        label_visibility="collapsed"
                    )
                    if new_text != q_text:
                        update_question(category_index, i, new_text)
                with col2:
                    if st.button("📝", key=f"instruction_{category_index}_{i}", help="Add instructions"):
                        st.session_state[f'show_instruction_{category_index}_{i}'] = True
                with col3:
                    if st.button("🗑️", key=f"delete_{category_index}_{i}", help="Delete question"):
                        delete_question(category_index, i)
                        st.rerun()
                
                if st.session_state.get(f'show_instruction_{category_index}_{i}', False):
                    with st.expander("Question Instructions", expanded=True):
                        instruction = st.text_area(
                            "Instructions",
                            value=question["instruction"] if isinstance(question, dict) and "instruction" in question else "",
                            key=f"instruction_text_{category_index}_{i}",
                            height=100,
                            label_visibility="visible"
                        )
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("Save", key=f"save_instruction_{category_index}_{i}"):
                                save_instruction(category_index, i, instruction)
                                st.success("Instructions saved!")
                                st.session_state[f'show_instruction_{category_index}_{i}'] = False
                        with col2:
                            if st.button("Cancel", key=f"cancel_instruction_{category_index}_{i}"):
                                st.session_state[f'show_instruction_{category_index}_{i}'] = False
        
        # Add question button at the end of each category
        if st.button("Add Question", key=f"add_question_{category_index}"):
            add_question(category_index)
            st.rerun()
        
        st.markdown("---")  # Add a separator after each category

def update_question(category_index, question_index, new_text):
    st.session_state['questions']["project_questions"][category_index]["questions"][question_index]["text"] = new_text
    save_questions_to_file(st.session_state['questions'])

def add_question(category_index):
    new_question = {"text": "New question", "instruction": ""}
    st.session_state['questions']["project_questions"][category_index]["questions"].append(new_question)
    save_questions_to_file(st.session_state['questions'])

def delete_question(category_index, question_index):
    del st.session_state['questions']["project_questions"][category_index]["questions"][question_index]
    save_questions_to_file(st.session_state['questions'])

def save_instruction(category_index, question_index, instruction):
    questions = st.session_state['questions']["project_questions"][category_index]["questions"]
    if isinstance(questions[question_index], str):
        questions[question_index] = {
            "text": questions[question_index],
            "instruction": instruction
        }
    else:
        questions[question_index]["instruction"] = instruction
    
    save_questions_to_file(st.session_state['questions'])

def save_questions_to_file(questions):
    with open('Questions.json', 'w') as f:
        json.dump(questions, f, indent=2)

def load_questions():
    try:
        with open('Questions.json', 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        st.error("Questions.json file not found. Please ensure it exists in the same directory as the script.")
        return {"project_questions": []}

if __name__ == "__main__":
    main()